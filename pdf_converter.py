#!/usr/bin/env python3

import io
import os
import shutil
import sys
from pathlib import Path
from typing import List, Optional, Tuple
import argparse

import re

try:
    import pdfplumber
    import pytesseract
    from pypdf import PdfReader, PdfWriter
    from pdf2image import convert_from_path
    from PIL import Image, ImageEnhance
    import docx
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
except ImportError as e:
    print(f"Missing required package: {e}")
    print("Install with: pip install pypdf pdfplumber pytesseract pdf2image pillow python-docx")
    sys.exit(1)


class PDFConverter:
    """Main class for converting PDFs to readable formats"""
    
    def __init__(self, input_pdf: str, output_format: str = 'pdf'):
        self.input_pdf = Path(input_pdf)
        self.output_format = output_format.lower()
        
        if not self.input_pdf.exists():
            raise FileNotFoundError(f"PDF file not found: {input_pdf}")
        
        if self.output_format not in ['pdf', 'txt', 'docx']:
            raise ValueError(f"Unsupported output format: {output_format}. Use: pdf, txt, docx")
    
    def is_scanned_pdf(self) -> bool:
        """
        Determine if PDF is scanned (image-based) or digital (text-based)
        Returns True if scanned, False if digital
        """
        try:
            with pdfplumber.open(self.input_pdf) as pdf:
                # Check first 3 pages for text
                pages_to_check = min(3, len(pdf.pages))
                total_text_length = 0
                
                for i in range(pages_to_check):
                    text = pdf.pages[i].extract_text()
                    if text:
                        total_text_length += len(text.strip())
                
                # Handle empty PDF
                if pages_to_check == 0:
                    return True  # Treat empty PDF as scanned (OCR will yield nothing)

                # If very little text found, likely scanned
                # Threshold: less than 50 characters per page on average
                avg_text_per_page = total_text_length / pages_to_check
                return avg_text_per_page < 50
        except Exception as e:
            print(f"Warning: Could not determine PDF type: {e}")
            return True  # Assume scanned if unsure
    
    def extract_text_digital(self) -> str:
        """Extract text from digital (non-scanned) PDF"""
        print("Extracting text from digital PDF...")
        text = []
        
        with pdfplumber.open(self.input_pdf) as pdf:
            for i, page in enumerate(pdf.pages, 1):
                print(f"Processing page {i}/{len(pdf.pages)}...")
                page_text = page.extract_text()
                
                if page_text:
                    text.append(f"--- Page {i} ---\n\n{page_text}\n\n")
                else:
                    text.append(f"--- Page {i} ---\n\n[No text found on this page]\n\n")
        
        return ''.join(text)

    def create_searchable_pdf(self, dpi: int = 350, lang: str = 'eng',
                              preprocess: bool = True,
                              preserve_color: bool = True,
                              output_path: Optional[str] = None) -> str:
        """
        Create a searchable PDF by adding an invisible text layer via OCR.
        Returns the same PDF with selectable/copyable text.
        """
        if output_path is None:
            output_path = self.input_pdf.with_stem(
                self.input_pdf.stem + '_searchable'
            ).with_suffix('.pdf')
        output_path = Path(output_path)

        # Verify Tesseract is available
        try:
            pytesseract.get_tesseract_version()
        except pytesseract.TesseractNotFoundError:
            print("Error: Tesseract OCR is not installed or not in PATH.")
            print("Install from: https://github.com/tesseract-ocr/tesseract")
            raise

        print(f"Converting PDF to images (DPI: {dpi})...")
        try:
            images = convert_from_path(
                self.input_pdf,
                dpi=dpi,
                fmt='png',  # PNG preserves quality better for OCR
                thread_count=4
            )
        except Exception as e:
            print(f"Error converting PDF to images: {e}")
            if sys.platform == 'win32':
                print("On Windows: Install poppler from https://github.com/osber/poppler-windows/releases")
            else:
                print("On Linux: apt-get install poppler-utils")
            raise

        print(f"Creating searchable PDF from {len(images)} pages...")
        pdf_writer = PdfWriter()

        for i, image in enumerate(images, 1):
            print(f"OCR progress: {i}/{len(images)}...")
            try:
                if preprocess:
                    image = self._preprocess_image_for_ocr(image, preserve_color=preserve_color)
                # Tesseract creates a single-page PDF with image + invisible text layer
                page_pdf_bytes = pytesseract.image_to_pdf_or_hocr(
                    image,
                    lang=lang,
                    config='--psm 3',
                    extension='pdf'
                )
                page_reader = PdfReader(io.BytesIO(page_pdf_bytes))
                pdf_writer.add_page(page_reader.pages[0])
            except Exception as e:
                print(f"Warning: OCR failed on page {i}, retrying without preprocessing...")
                try:
                    orig_image = images[i - 1]
                    page_pdf_bytes = pytesseract.image_to_pdf_or_hocr(
                        orig_image, lang=lang, config='--psm 3', extension='pdf'
                    )
                    page_reader = PdfReader(io.BytesIO(page_pdf_bytes))
                    pdf_writer.add_page(page_reader.pages[0])
                except Exception as e2:
                    raise RuntimeError(f"Could not process page {i}: {e2}") from e

        with open(output_path, 'wb') as f:
            pdf_writer.write(f)

        print(f"Saved searchable PDF to: {output_path}")
        return str(output_path)

    def _preprocess_image_for_ocr(self, image: Image.Image,
                                   preserve_color: bool = False) -> Image.Image:
        """
        Preprocess image to improve OCR accuracy: contrast, sharpen.
        If preserve_color is False, converts to grayscale (better OCR, smaller output).
        If preserve_color is True, keeps original colors in output.
        """
        if not preserve_color and image.mode != 'L':
            image = image.convert('L')
        # Enhance contrast (helps with faded scans)
        enhancer = ImageEnhance.Contrast(image)
        image = enhancer.enhance(1.5)
        # Sharpen (reduces blur from scanning)
        enhancer = ImageEnhance.Sharpness(image)
        image = enhancer.enhance(1.5)
        return image

    def extract_text_ocr(self, dpi: int = 300, lang: str = 'eng',
                         preprocess: bool = True) -> str:
        """
        Extract text from scanned PDF using OCR
        
        Args:
            dpi: Resolution for image conversion (higher = better quality, slower)
            lang: Tesseract language code (eng, fra, deu, etc.)
        """
        print(f"Converting PDF to images (DPI: {dpi})...")
        print("This may take a while for large documents...")

        # Verify Tesseract is available before processing
        try:
            pytesseract.get_tesseract_version()
        except pytesseract.TesseractNotFoundError:
            print("Error: Tesseract OCR is not installed or not in PATH.")
            print("Install from: https://github.com/tesseract-ocr/tesseract")
            raise

        try:
            # Convert PDF to images
            images = convert_from_path(
                self.input_pdf,
                dpi=dpi,
                fmt='jpeg',
                thread_count=4
            )
        except Exception as e:
            print(f"Error converting PDF to images: {e}")
            if sys.platform == 'win32':
                print("On Windows: Install poppler from https://github.com/osber/poppler-windows/releases")
                print("Or: conda install -c conda-forge poppler")
            else:
                print("On Linux: apt-get install poppler-utils")
            raise
        
        print(f"Performing OCR on {len(images)} pages...")
        text = []
        
        for i, image in enumerate(images, 1):
            print(f"OCR progress: {i}/{len(images)}...")

            try:
                # Preprocess image for better OCR accuracy
                if preprocess:
                    image = self._preprocess_image_for_ocr(image)
                # Perform OCR (PSM 3 = fully automatic page segmentation, best for documents)
                page_text = pytesseract.image_to_string(
                    image,
                    lang=lang,
                    config='--psm 3'
                )
                text.append(f"--- Page {i} ---\n\n{page_text}\n\n")
            except Exception as e:
                print(f"Warning: OCR failed on page {i}: {e}")
                text.append(f"--- Page {i} ---\n\n[OCR failed for this page]\n\n")
        
        return ''.join(text)
    
    def save_as_txt(self, text: str, output_path: Optional[str] = None) -> str:
        """Save extracted text as TXT file"""
        if output_path is None:
            output_path = self.input_pdf.with_suffix('.txt')
        
        output_path = Path(output_path)
        # Handle encoding issues from OCR (replace unencodable chars)
        safe_text = text.encode('utf-8', errors='replace').decode('utf-8')
        output_path.write_text(safe_text, encoding='utf-8')
        print(f"Saved text to: {output_path}")
        return str(output_path)
    
    def save_as_docx(self, text: str, output_path: Optional[str] = None) -> str:
        """Save extracted text as DOCX file with basic formatting"""
        if output_path is None:
            output_path = self.input_pdf.with_suffix('.docx')
        
        output_path = Path(output_path)
        
        # Create Word document
        doc = docx.Document()
        
        # Add title
        title = doc.add_paragraph()
        title_run = title.add_run(f"Extracted from: {self.input_pdf.name}")
        title_run.bold = True
        title_run.font.size = Pt(14)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.add_paragraph()  # Empty line

        # Split text by pages using regex (avoids conflicts if content contains "--- Page ---")
        page_pattern = re.compile(r'--- Page (\d+) ---\n\n', re.IGNORECASE)
        parts = page_pattern.split(text)
        # parts[0] = leading text (empty), then page_num, content, page_num, content, ...
        for i in range(1, len(parts), 2):
            if i + 1 < len(parts):
                page_num = parts[i]
                content = parts[i + 1].strip()
            else:
                continue

            # Add page heading
            heading = doc.add_paragraph()
            heading_run = heading.add_run(f"Page {page_num}")
            heading_run.bold = True
            heading_run.font.size = Pt(12)
            try:
                heading_run.font.color.rgb = RGBColor(0, 0, 139)
            except (AttributeError, KeyError):
                pass  # Some python-docx versions handle color differently

            # Add content
            doc.add_paragraph(content)
            doc.add_paragraph()  # Space between pages
        
        doc.save(str(output_path))
        print(f"Saved Word document to: {output_path}")
        return str(output_path)
    
    def convert(self, force_ocr: bool = False, dpi: int = 350,
                lang: str = 'eng', output_path: Optional[str] = None,
                preprocess: bool = True, preserve_color: bool = True) -> str:
        """
        Main conversion method
        
        Args:
            force_ocr: Force OCR even if PDF appears to have text
            dpi: DPI for image conversion (if OCR is needed)
            lang: Language for OCR
            output_path: Custom output path
        
        Returns:
            Path to the output file
        """
        print(f"Converting: {self.input_pdf}")
        print(f"Output format: {self.output_format.upper()}")
        print()
        
        # Determine if OCR is needed
        if force_ocr:
            print("Force OCR enabled - using OCR regardless of PDF type")
            use_ocr = True
        else:
            use_ocr = self.is_scanned_pdf()
            if use_ocr:
                print("Detected scanned PDF - will use OCR")
            else:
                print("Detected digital PDF - extracting text directly")
        
        print()

        # Output as searchable PDF (primary use case)
        if self.output_format == 'pdf':
            if use_ocr:
                return self.create_searchable_pdf(
                    dpi=dpi, lang=lang, preprocess=preprocess,
                    preserve_color=preserve_color, output_path=output_path
                )
            else:
                # Digital PDF is already searchable - copy to output
                out = Path(output_path) if output_path else self.input_pdf.with_stem(
                    self.input_pdf.stem + '_searchable'
                ).with_suffix('.pdf')
                shutil.copy2(self.input_pdf, out)
                print(f"PDF already has text layer - copied to: {out}")
                return str(out)

        # Extract text for TXT/DOCX output
        if use_ocr:
            text = self.extract_text_ocr(dpi=dpi, lang=lang, preprocess=preprocess)
        else:
            text = self.extract_text_digital()

        print()
        print(f"Extracted {len(text)} characters")
        print()

        if self.output_format == 'txt':
            return self.save_as_txt(text, output_path)
        elif self.output_format == 'docx':
            return self.save_as_docx(text, output_path)


def main():
    parser = argparse.ArgumentParser(
        description='Convert scanned PDFs to searchable PDFs (adds text layer via OCR)',
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog="""
Examples:
  # Create searchable PDF (default - returns PDF with selectable text)
  python pdf_converter.py document.pdf
  
  # Specify output file
  python pdf_converter.py document.pdf -o searchable.pdf
  
  # Extract to text or Word instead
  python pdf_converter.py document.pdf --format txt
  python pdf_converter.py document.pdf --format docx
  
  # Force OCR with custom DPI
  python pdf_converter.py document.pdf --force-ocr --dpi 400
  
  # OCR with different language (e.g., French)
  python pdf_converter.py document.pdf --lang fra
        """
    )
    
    parser.add_argument('input_pdf', help='Input PDF file to convert')
    parser.add_argument('-f', '--format', choices=['pdf', 'txt', 'docx'],
                       default='pdf', help='Output format (default: pdf)')
    parser.add_argument('-o', '--output', help='Output file path (optional)')
    parser.add_argument('--force-ocr', action='store_true',
                       help='Force OCR even if PDF has selectable text')
    parser.add_argument('--dpi', type=int, default=350,
                       help='DPI for image conversion when using OCR (default: 350)')
    parser.add_argument('--lang', default='eng',
                       help='Tesseract language code for OCR (default: eng)')
    parser.add_argument('--grayscale', action='store_true',
                       help='Output grayscale PDF (default: preserve color)')
    
    args = parser.parse_args()
    
    try:
        converter = PDFConverter(args.input_pdf, args.format)
        output_file = converter.convert(
            force_ocr=args.force_ocr,
            dpi=args.dpi,
            lang=args.lang,
            output_path=args.output,
            preserve_color=not args.grayscale
        )
        print()
        print("=" * 60)
        print("SUCCESS!")
        print(f"Output file: {output_file}")
        print("=" * 60)
        
    except Exception as e:
        print()
        print("=" * 60)
        print("ERROR!")
        print(f"{type(e).__name__}: {e}")
        print("=" * 60)
        sys.exit(1)


if __name__ == '__main__':
    main()
