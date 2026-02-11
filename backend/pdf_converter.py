#!/usr/bin/env python3
"""
PDF to Readable Format Converter
Converts scanned PDFs into searchable PDFs by adding an invisible text layer via OCR.
Returns the same PDF file with selectable/copyable text. Also supports TXT and DOCX output.
"""

import io
import os
import shutil
import sys
from pathlib import Path
from typing import List, Optional, Tuple
import argparse
import re

try:
    import pdfplumber
    import pytesseract
    from pypdf import PdfReader, PdfWriter
    from pdf2image import convert_from_path
    from PIL import Image, ImageEnhance
    import docx
    from docx.shared import Pt, RGBColor
except ImportError as e:
    print(f"Missing required package: {e}")
    print("Install with: pip install pypdf pdfplumber pytesseract pdf2image pillow python-docx")
    sys.exit(1)


class PDFConverter:
    """Main class for converting PDFs to readable formats"""

    def __init__(self, input_pdf: str, output_format: str = 'pdf'):
        self.input_pdf = Path(input_pdf)
        self.output_format = output_format.lower()

        if not self.input_pdf.exists():
            raise FileNotFoundError(f"PDF file not found: {input_pdf}")

        if self.output_format not in ['pdf', 'txt', 'docx']:
            raise ValueError(f"Unsupported output format: {output_format}. Use: pdf, txt, docx")

    def is_scanned_pdf(self) -> bool:
        """Determine if PDF is scanned (image-based) or digital (text-based)"""
        try:
            with pdfplumber.open(self.input_pdf) as pdf:
                pages_to_check = min(3, len(pdf.pages))
                total_text_length = 0

                for i in range(pages_to_check):
                    text = pdf.pages[i].extract_text()
                    if text:
                        total_text_length += len(text.strip())

                if pages_to_check == 0:
                    return True

                avg_text_per_page = total_text_length / pages_to_check
                return avg_text_per_page < 50
        except Exception as e:
            print(f"Warning: Could not determine PDF type: {e}")
            return True

    def extract_text_digital(self) -> str:
        """Extract text from digital (non-scanned) PDF"""
        print("Extracting text from digital PDF...")
        text = []
        with pdfplumber.open(self.input_pdf) as pdf:
            for i, page in enumerate(pdf.pages, 1):
                print(f"Processing page {i}/{len(pdf.pages)}...")
                page_text = page.extract_text()
                if page_text:
                    text.append(f"--- Page {i} ---\n\n{page_text}\n\n")
                else:
                    text.append(f"--- Page {i} ---\n\n[No text found on this page]\n\n")
        return ''.join(text)

    def create_searchable_pdf(self, dpi: int = 350, lang: str = 'eng',
                              preprocess: bool = True,
                              preserve_color: bool = True,
                              output_path: Optional[str] = None) -> str:
        """Create a searchable PDF by adding an invisible text layer via OCR."""
        if output_path is None:
            output_path = self.input_pdf.with_stem(
                self.input_pdf.stem + '_searchable'
            ).with_suffix('.pdf')
        output_path = Path(output_path)

        try:
            pytesseract.get_tesseract_version()
        except pytesseract.TesseractNotFoundError:
            print("Error: Tesseract OCR is not installed or not in PATH.")
            raise

        print(f"Converting PDF to images (DPI: {dpi})...")
        try:
            images = convert_from_path(
                self.input_pdf, dpi=dpi, fmt='png',
                thread_count=4
            )
        except Exception as e:
            print(f"Error converting PDF to images: {e}")
            raise

        print(f"Creating searchable PDF from {len(images)} pages...")
        pdf_writer = PdfWriter()

        for i, image in enumerate(images, 1):
            print(f"OCR progress: {i}/{len(images)}...")
            try:
                if preprocess:
                    image = self._preprocess_image_for_ocr(image, preserve_color=preserve_color)
                page_pdf_bytes = pytesseract.image_to_pdf_or_hocr(
                    image, lang=lang, config='--psm 3', extension='pdf'
                )
                page_reader = PdfReader(io.BytesIO(page_pdf_bytes))
                pdf_writer.add_page(page_reader.pages[0])
            except Exception as e:
                print(f"Warning: OCR failed on page {i}, retrying without preprocessing...")
                try:
                    orig_image = images[i - 1]
                    page_pdf_bytes = pytesseract.image_to_pdf_or_hocr(
                        orig_image, lang=lang, config='--psm 3', extension='pdf'
                    )
                    page_reader = PdfReader(io.BytesIO(page_pdf_bytes))
                    pdf_writer.add_page(page_reader.pages[0])
                except Exception as e2:
                    raise RuntimeError(f"Could not process page {i}: {e2}") from e

        with open(output_path, 'wb') as f:
            pdf_writer.write(f)

        print(f"Saved searchable PDF to: {output_path}")
        return str(output_path)

    def _preprocess_image_for_ocr(self, image: Image.Image,
                                   preserve_color: bool = False) -> Image.Image:
        """Preprocess image: contrast, sharpen. Optionally convert to grayscale."""
        if not preserve_color and image.mode != 'L':
            image = image.convert('L')
        enhancer = ImageEnhance.Contrast(image)
        image = enhancer.enhance(1.5)
        enhancer = ImageEnhance.Sharpness(image)
        image = enhancer.enhance(1.5)
        return image

    def extract_text_ocr(self, dpi: int = 300, lang: str = 'eng',
                         preprocess: bool = True) -> str:
        """Extract text from scanned PDF using OCR"""
        print(f"Converting PDF to images (DPI: {dpi})...")
        try:
            pytesseract.get_tesseract_version()
        except pytesseract.TesseractNotFoundError:
            raise

        try:
            images = convert_from_path(
                self.input_pdf, dpi=dpi, fmt='jpeg', thread_count=4
            )
        except Exception as e:
            print(f"Error converting PDF to images: {e}")
            raise

        print(f"Performing OCR on {len(images)} pages...")
        text = []
        for i, image in enumerate(images, 1):
            print(f"OCR progress: {i}/{len(images)}...")
            try:
                if preprocess:
                    image = self._preprocess_image_for_ocr(image)
                page_text = pytesseract.image_to_string(
                    image, lang=lang, config='--psm 3'
                )
                text.append(f"--- Page {i} ---\n\n{page_text}\n\n")
            except Exception as e:
                print(f"Warning: OCR failed on page {i}: {e}")
                text.append(f"--- Page {i} ---\n\n[OCR failed for this page]\n\n")
        return ''.join(text)

    def save_as_txt(self, text: str, output_path: Optional[str] = None) -> str:
        """Save extracted text as TXT file"""
        if output_path is None:
            output_path = self.input_pdf.with_suffix('.txt')
        output_path = Path(output_path)
        safe_text = text.encode('utf-8', errors='replace').decode('utf-8')
        output_path.write_text(safe_text, encoding='utf-8')
        return str(output_path)

    def save_as_docx(self, text: str, output_path: Optional[str] = None) -> str:
        """Save extracted text as DOCX file. No alignment is set so you can edit freely."""
        if output_path is None:
            output_path = self.input_pdf.with_suffix('.docx')
        output_path = Path(output_path)
        doc = docx.Document()
        title = doc.add_paragraph()
        title_run = title.add_run(f"Extracted from: {self.input_pdf.name}")
        title_run.bold = True
        title_run.font.size = Pt(14)
        # No alignment set — leave default so you can edit the document freely
        doc.add_paragraph()

        page_pattern = re.compile(r'--- Page (\d+) ---\n\n', re.IGNORECASE)
        parts = page_pattern.split(text)
        for i in range(1, len(parts), 2):
            if i + 1 < len(parts):
                page_num = parts[i]
                content = parts[i + 1].strip()
            else:
                continue
            heading = doc.add_paragraph()
            heading_run = heading.add_run(f"Page {page_num}")
            heading_run.bold = True
            heading_run.font.size = Pt(12)
            try:
                heading_run.font.color.rgb = RGBColor(0, 0, 139)
            except (AttributeError, KeyError):
                pass
            doc.add_paragraph(content)
            doc.add_paragraph()

        doc.save(str(output_path))
        return str(output_path)

    def convert(self, force_ocr: bool = False, dpi: int = 350,
                lang: str = 'eng', output_path: Optional[str] = None,
                preprocess: bool = True, preserve_color: bool = True) -> str:
        """Main conversion method"""
        print(f"Converting: {self.input_pdf}")
        use_ocr = force_ocr or self.is_scanned_pdf()

        if self.output_format == 'pdf':
            if use_ocr:
                return self.create_searchable_pdf(
                    dpi=dpi, lang=lang, preprocess=preprocess,
                    preserve_color=preserve_color, output_path=output_path
                )
            else:
                out = Path(output_path) if output_path else self.input_pdf.with_stem(
                    self.input_pdf.stem + '_searchable'
                ).with_suffix('.pdf')
                shutil.copy2(self.input_pdf, out)
                return str(out)

        if use_ocr:
            text = self.extract_text_ocr(dpi=dpi, lang=lang, preprocess=preprocess)
        else:
            text = self.extract_text_digital()

        if self.output_format == 'txt':
            return self.save_as_txt(text, output_path)
        elif self.output_format == 'docx':
            return self.save_as_docx(text, output_path)

        return str(output_path)


def main():
    """CLI entry point"""
    parser = argparse.ArgumentParser(description='Convert scanned PDFs to searchable PDFs')
    parser.add_argument('input_pdf', help='Input PDF file')
    parser.add_argument('-f', '--format', choices=['pdf', 'txt', 'docx'], default='pdf')
    parser.add_argument('-o', '--output', help='Output file path')
    parser.add_argument('--force-ocr', action='store_true')
    parser.add_argument('--dpi', type=int, default=350)
    parser.add_argument('--lang', default='eng')
    parser.add_argument('--grayscale', action='store_true')

    args = parser.parse_args()
    try:
        converter = PDFConverter(args.input_pdf, args.format)
        output_file = converter.convert(
            force_ocr=args.force_ocr, dpi=args.dpi, lang=args.lang,
            output_path=args.output, preserve_color=not args.grayscale
        )
        print(f"SUCCESS! Output: {output_file}")
    except Exception as e:
        print(f"ERROR: {e}")
        sys.exit(1)


if __name__ == '__main__':
    main()
