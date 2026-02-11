"""
DOCX <-> HTML conversion for the in-browser editor.
- docx_to_html: for loading DOCX into the editor (mammoth).
- html_to_docx: for saving editor content back to DOCX (no alignment set).
"""

import re
from pathlib import Path
from typing import Optional

try:
    import mammoth
    from bs4 import BeautifulSoup
    import docx
    from docx.shared import Pt, RGBColor
except ImportError as e:
    raise ImportError(
        "Install: pip install mammoth beautifulsoup4 python-docx"
    ) from e


def docx_to_html(docx_path: Path) -> str:
    """Convert a DOCX file to HTML for the editor. Whole document, as-is: structure and images preserved."""
    path = Path(docx_path)
    if not path.exists() or path.suffix.lower() != ".docx":
        raise ValueError("Not a DOCX file or file not found")
    with open(path, "rb") as f:
        result = mammoth.convert_to_html(
            f,
            ignore_empty_paragraphs=False,  # keep all paragraphs and spacing so document opens as-is
            convert_image=mammoth.images.data_uri,  # embed images inline so they show in editor
        )
    return result.value or "<p></p>"


def _parse_font_size(style: str) -> Optional[float]:
    """Parse font-size from CSS style string. Returns size in points or None."""
    if not style:
        return None
    # font-size: 12pt; or font-size: 14px; etc.
    match = re.search(r"font-size:\s*([\d.]+)\s*(pt|px|em)?", style, re.I)
    if not match:
        return None
    try:
        value = float(match.group(1))
    except ValueError:
        return None
    unit = (match.group(2) or "pt").lower()
    if unit == "pt":
        return value
    if unit == "px":
        return value * 0.75  # approximate
    if unit == "em":
        return value * 12
    return value


def _parse_color(style: str) -> Optional[tuple]:
    """Parse color from CSS style. Returns (r, g, b) 0-255 or None."""
    if not style:
        return None
    # color: rgb(0, 0, 139); or color: #00008b;
    match = re.search(r"color:\s*rgb\s*\(\s*(\d+)\s*,\s*(\d+)\s*,\s*(\d+)\s*\)", style, re.I)
    if match:
        return (int(match.group(1)), int(match.group(2)), int(match.group(3)))
    match = re.search(r"color:\s*#([0-9a-fA-F]{6})", style)
    if match:
        h = match.group(1)
        return (int(h[0:2], 16), int(h[2:4], 16), int(h[4:6], 16))
    return None


def _add_inline_format(run, tag, style: str):
    """Apply inline formatting to a docx run from tag name and optional style."""
    tag_lower = tag.name.lower() if hasattr(tag, "name") else ""
    if tag_lower in ("strong", "b"):
        run.bold = True
    elif tag_lower in ("em", "i"):
        run.italic = True
    elif tag_lower == "u":
        run.underline = True
    if style:
        pt = _parse_font_size(style)
        if pt is not None:
            run.font.size = Pt(min(max(pt, 6), 72))
        rgb = _parse_color(style)
        if rgb is not None:
            run.font.color.rgb = RGBColor(*rgb)


def _process_inline(element, paragraph):
    """Recursively process inline elements and text into a docx paragraph."""
    if element is None:
        return
    if isinstance(element, str):
        text = element.strip()
        if text:
            run = paragraph.add_run(text)
        return
    if not hasattr(element, "name"):
        return
    tag = element
    name = tag.name.lower()

    if name == "br":
        # Add line break within paragraph
        run = paragraph.add_run()
        run.add_br()
        return

    style = tag.get("style", "") or ""

    if name in ("strong", "b", "em", "i", "u", "span"):
        for child in tag.children:
            if child.name:
                _process_inline(child, paragraph)
            else:
                text = str(child).strip()
                if text:
                    run = paragraph.add_run(text)
                    _add_inline_format(run, tag, style)
        return

    # For other inlines (e.g. a), just add text
    if tag.string:
        run = paragraph.add_run(tag.string.strip())
        _add_inline_format(run, tag, style)
    else:
        for child in tag.children:
            _process_inline(child, paragraph)


def _is_block(tag) -> bool:
    if not hasattr(tag, "name"):
        return False
    return tag.name.lower() in (
        "p", "div", "h1", "h2", "h3", "h4", "h5", "h6",
        "ul", "ol", "li", "blockquote", "hr"
    )


def html_to_docx(html: str, output_path: Path) -> Path:
    """
    Convert HTML from the editor to DOCX. Does not set paragraph alignment
    so the document remains editable without alignment overrides.
    """
    soup = BeautifulSoup(html, "html.parser")
    doc = docx.Document()

    # Normalize: wrap raw text in body if needed
    body = soup.find("body") or soup
    if body.name != "body":
        # Maybe the whole thing is a fragment
        blocks = list(body.children) if hasattr(body, "children") else [body]
    else:
        blocks = list(body.children)

    for node in blocks:
        if not hasattr(node, "name"):
            text = str(node).strip() if isinstance(node, str) else ""
            if text:
                p = doc.add_paragraph()
                p.add_run(text)
            continue
        tag = node
        name = tag.name.lower() if tag.name else ""

        if name in ("p", "div"):
            p = doc.add_paragraph()
            for child in tag.children:
                _process_inline(child, p)
        elif name in ("h1", "h2", "h3", "h4", "h5", "h6"):
            level = int(name[1])
            p = doc.add_paragraph()
            p.style = f"Heading {level}"
            for child in tag.children:
                _process_inline(child, p)
        elif name == "br":
            doc.add_paragraph()
        elif name == "hr":
            doc.add_paragraph()
        elif name in ("ul", "ol"):
            for li in tag.find_all("li", recursive=False):
                p = doc.add_paragraph(style="List Bullet" if name == "ul" else "List Number")
                for child in li.children:
                    _process_inline(child, p)
        elif name == "li":
            p = doc.add_paragraph(style="List Bullet")
            for child in tag.children:
                _process_inline(child, p)
        else:
            # Unknown block: treat as paragraph
            p = doc.add_paragraph()
            for child in tag.children:
                _process_inline(child, p)

    if len(doc.paragraphs) == 0:
        doc.add_paragraph()
    out = Path(output_path)
    doc.save(str(out))
    return out
